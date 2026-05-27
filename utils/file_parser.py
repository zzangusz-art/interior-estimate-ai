# -*- coding: utf-8 -*-
"""
파일 파싱 유틸리티
지원 형식: PDF, Excel(.xlsx/.xls), Word(.docx), 이미지(PNG/JPG)
"""
import io
import os
import base64
from pathlib import Path


def parse_file(file_path: str, filename: str) -> dict:
    """파일을 파싱해서 텍스트/테이블/이미지 데이터를 반환"""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        return _parse_image(file_path)
    elif ext == ".hwp":
        return _parse_hwp(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")


def _parse_pdf(file_path: str) -> dict:
    """PDF 파싱: 텍스트 추출을 시도하고, 깨진 경우 이미지 변환으로 폴백"""
    import pdfplumber

    text_parts = []
    tables = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(f"[페이지 {i+1}]\n{text}")
            page_tables = page.extract_tables()
            for tbl in page_tables:
                if tbl:
                    tables.append({"page": i + 1, "data": tbl})

    full_text = "\n\n".join(text_parts)

    # 텍스트가 깨진 경우(비율 기반 휴리스틱) → 이미지로 변환
    if _is_garbled(full_text):
        return _pdf_to_image(file_path)

    return {
        "type": "pdf",
        "text": full_text,
        "tables": tables,
        "image_base64": None,
    }


def _is_garbled(text: str) -> bool:
    """텍스트가 인코딩 오류로 깨진 경우 감지"""
    if not text or len(text) < 20:
        return True
    # 대체 문자(U+FFFD) 또는 제어 문자 비율이 높으면 깨진 것으로 판단
    garbage = sum(1 for c in text if ord(c) > 0xAC00 + 11172 or c in "�\x00")
    total_non_space = sum(1 for c in text if not c.isspace())
    if total_non_space == 0:
        return True
    ratio = garbage / total_non_space
    return ratio > 0.15


def _pdf_to_image(file_path: str) -> dict:
    """PDF를 이미지로 변환하여 Vision API용 데이터 반환 (첫 3페이지)"""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(file_path)
    page_images = []

    max_pages = min(len(pdf), 3)
    for i in range(max_pages):
        page = pdf[i]
        bitmap = page.render(scale=2.0)  # 144dpi
        pil_image = bitmap.to_pil()

        buf = io.BytesIO()
        pil_image.save(buf, format="PNG")
        page_images.append(buf.getvalue())

    if not page_images:
        return {"type": "pdf", "text": "", "tables": [], "image_base64": None}

    # 페이지가 여러 장이면 세로로 합치기
    if len(page_images) == 1:
        final_bytes = page_images[0]
    else:
        from PIL import Image as PILImage
        imgs = [PILImage.open(io.BytesIO(b)) for b in page_images]
        total_height = sum(img.height for img in imgs)
        max_width = max(img.width for img in imgs)
        combined = PILImage.new("RGB", (max_width, total_height), (255, 255, 255))
        y_offset = 0
        for img in imgs:
            combined.paste(img, (0, y_offset))
            y_offset += img.height
        buf = io.BytesIO()
        combined.save(buf, format="PNG")
        final_bytes = buf.getvalue()

    b64 = base64.standard_b64encode(final_bytes).decode("utf-8")
    return {
        "type": "pdf_image",
        "text": "",
        "tables": [],
        "image_base64": b64,
        "media_type": "image/png",
    }


def _parse_excel(file_path: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    text_parts = []
    tables = []

    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append([str(c) if c is not None else "" for c in row])
        if rows:
            tables.append({"sheet": sheet.title, "data": rows})
            text_parts.append(f"[시트: {sheet.title}]")
            for row in rows:
                text_parts.append("\t".join(row))

    return {
        "type": "excel",
        "text": "\n".join(text_parts),
        "tables": tables,
        "image_base64": None,
    }


def _parse_docx(file_path: str) -> dict:
    from docx import Document
    doc = Document(file_path)
    text_parts = []
    tables = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"data": rows})
        for row in rows:
            text_parts.append("\t".join(row))

    return {
        "type": "docx",
        "text": "\n".join(text_parts),
        "tables": tables,
        "image_base64": None,
    }


def _parse_image(file_path: str) -> dict:
    """
    이미지 파일을 Vision API용으로 최적화하여 반환.
    - EXIF 회전 자동 보정 (폰 카메라 사진)
    - 최장 변 2048px 이하로 리사이즈
    - JPEG quality 조정으로 5MB 이내 보장
    """
    from PIL import Image as PILImage, ImageOps

    # 파일 크기 체크 (20MB 초과 거부)
    file_size = os.path.getsize(file_path)
    if file_size > 20 * 1024 * 1024:
        raise ValueError("이미지 파일이 너무 큽니다 (최대 20MB)")

    with PILImage.open(file_path) as img:
        # EXIF 회전 보정 (iPhone/Android 카메라 사진 대응)
        img = ImageOps.exif_transpose(img)

        # 투명도 채널 제거 (RGBA → RGB, P → RGB)
        if img.mode in ("RGBA", "P", "LA"):
            background = PILImage.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")

        # 최장 변 2048px 이하로 리사이즈 (과도하게 큰 사진 처리)
        MAX_DIM = 2048
        w, h = img.size
        if max(w, h) > MAX_DIM:
            ratio = MAX_DIM / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), PILImage.LANCZOS)

        # JPEG로 인코딩 (quality 단계적 조정으로 4MB 이내 보장)
        img_bytes = _encode_jpeg(img, target_max_bytes=4 * 1024 * 1024)

    b64 = base64.standard_b64encode(img_bytes).decode("utf-8")
    return {
        "type": "image",
        "text": "",
        "tables": [],
        "image_base64": b64,
        "media_type": "image/jpeg",
    }


def _encode_jpeg(img, target_max_bytes: int = 4 * 1024 * 1024) -> bytes:
    """목표 크기 이내로 JPEG 인코딩 (quality 단계적 감소)"""
    for quality in (92, 80, 65, 50):
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        data = buf.getvalue()
        if len(data) <= target_max_bytes:
            return data
    # 최후 수단: 절반 크기로 재시도
    w, h = img.size
    small = img.resize((w // 2, h // 2), PILImage.LANCZOS)
    buf = io.BytesIO()
    small.save(buf, format="JPEG", quality=75, optimize=True)
    return buf.getvalue()


def _parse_hwp(file_path: str) -> dict:
    """HWP: olefile로 텍스트 추출 시도, 실패 시 안내 메시지 반환"""
    try:
        import olefile
        ole = olefile.OleFileIO(file_path)
        if ole.exists("PrvText"):
            data = ole.openstream("PrvText").read()
            text = data.decode("utf-16-le", errors="ignore")
        else:
            text = "HWP 파일에서 텍스트를 추출할 수 없습니다. 내용을 직접 입력하거나 PDF/이미지로 변환 후 업로드해 주세요."
        return {"type": "hwp", "text": text, "tables": [], "image_base64": None}
    except Exception:
        return {
            "type": "hwp",
            "text": "HWP 파일 처리를 위해 olefile 라이브러리가 필요합니다. PDF 또는 이미지로 변환 후 업로드해 주세요.",
            "tables": [],
            "image_base64": None,
        }
